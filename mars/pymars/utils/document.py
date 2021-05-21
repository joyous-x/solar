# coding=utf-8

""" pdf-checker
pdf 相关工具的库主要有：at 2021-05-17
- [camelot](https://github.com/camelot-dev/camelot): 提取表格效果最好
    + 使用了 PyPDF2 来操作 pdf 文件
    + WARN：camelot\handlers.py: L73 打开文件后没有关闭，会句柄泄漏
- [pdfminer](https://github.com/pdfminer/pdfminer.six)
    + 可以解析 pdf 的树形结构，用于提取 文字、图片等
- [pdfplumber](https://github.com/jsvine/pdfplumber)
    + built on pdfminer.six

office 相关库：
- [python-docx](https://github.com/python-openxml/python-docx)
- [openpyxl](https://openpyxl.readthedocs.io/en/stable/)


# reference
#     https://openpyxl.readthedocs.io/en/stable/api/openpyxl.styles.fills.html
# note:
#     openpyxl 在插入一行时，如果后续有 mergedcell，则会错乱：不能处理mergedcell的格式
"""
import os
import io
import shutil
import collections

from .common import *

try:
    import camelot
    import PyPDF2
except Exception as e:
    camelot = None
    PyPDF2 = None

try:
    import PIL as pil
    import pdfplumber
except Exception as e:
    pdfplumber = None
    pil = None

from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage, PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator, TextConverter
from pdfminer.layout import LAParams, LTContainer,LTImage
from pdfminer.image import ImageWriter
import pdfminer.high_level
import pdfminer.layout

from openpyxl import load_workbook, Workbook
from openpyxl.worksheet.table import Table, TableStyleInfo
import zipfile
import docx


class CamelotWrapper(object):
    """ parse tables in pdf with Camelot
    reference: https://camelot-py.readthedocs.io/en/master/user/advanced.html
    """
    def __init__(self, filepath, passwd=None, pages="all", line_scale=40):
        """
        pages : str, optional (default: '1')
            Comma-separated page numbers. Example: '1,3,4' or '1,4-end' or 'all'.
        """
        self.filepath, self.passwd = filepath, passwd
        self.line_scale = line_scale
        self.pages = pages
        self.tables = None
        try:
            self.tables = camelot.read_pdf(filepath, pages="all" if not pages else pages, password=passwd, line_scale=line_scale)
        except Exception as e:
            print("!!! {} error: {}".format(filepath, e))

    def parse_tables(self, pages: set() = None, accuracy_min=10, debug=False):
        if self.tables is None:
            return None
        if debug:
            for table in self.tables:
                data_frame = table.df
                parsing_report = table.parsing_report
                print(parsing_report, "\n", data_frame)
        tables = list(filter(lambda t: t.parsing_report['accuracy'] > accuracy_min, self.tables))
        if pages is not None:
            tables = list(filter(lambda t: t.parsing_report['page'] in pages, tables))
        return tables

    @staticmethod
    def show(table, kind='grid'):
        camelot.plot(table, kind=kind).show()

    @staticmethod
    def pages_count(file_path, pdf_passwd=None):
        pages_count = None
        try:
            pdf_obj = PyPDF2.PdfFileReader(file_path, strict=False)
            pages_count = pdf_obj.getNumPages()
        except Exception as e:
            print("!!! {} error: {}".format(file_path, e))
        return pages_count

    @staticmethod
    def split_pages(file_path, page_set=None, output_dir=None, output_prefix="", pdf_passwd=None, debug=False):
        print("==> ready to split : {}".format(file_path))
        file_dir, file_name = os.path.split(file_path)
        file_name, file_suffix = os.path.splitext(file_name)
        if not output_dir:
            output_dir = file_dir
        pdf_obj = PyPDF2.PdfFileReader(file_path, strict=False)
        for pageno in range(1, pdf_obj.getNumPages()+1):
            if page_set and pageno not in page_set:
                continue
            pdf_writer = PyPDF2.PdfFileWriter()
            pdf_writer.addPage(pdf_obj.getPage(pageno-1))
            with open(os.path.join(output_dir, "{}{}.{}{}".format(output_prefix,file_name,pageno,file_suffix)), 'wb') as output_pdf:
                pdf_writer.write(output_pdf)
        print("===> end: total pages {}".format(pdf_obj.getNumPages()))
        return pdf_obj.getNumPages()

    @staticmethod
    def record_tables(tables, output_dir_path, prefix=""):
        if not tables:
            return None
        tables_cells_bound= list()
        for table in tables:
            table_cells_bound = list()
            for r in range(table.shape[0]):
                cells = [table.cells[r][c] for c in range(table.shape[1])]
                table_cells_bound.append([(cell.bottom, cell.top, cell.left, cell.right) for cell in cells])
            tables_cells_bound.append(table_cells_bound)
        for index, table in enumerate(tables_cells_bound):
            lines = [",".join([str(t) for t in line]) for line in table]
            with open(os.path.join(output_dir_path, "{}table_{}_bound.txt".format(prefix, index)), "a+") as f:
                f.write("\n".join(lines))
        for table in tables:
            lines = [" | ".join(line) for line in table.data]
            with open(os.path.join(output_dir_path, "{}table_{}_texts.txt".format(prefix, index)), "a+") as f:
                f.write("\n".join(lines))
        return len(tables)

class PdfminerWrapper(object):
    """
    reference: https://pdfminersix.readthedocs.io/en/latest/index.html
    """
    @staticmethod
    def flat_iter(obj):
        yield obj
        if isinstance(obj, LTContainer):
            for ob in obj:
                yield from PdfminerWrapper.flat_iter(ob)

    @staticmethod
    def pages_count(pdf_path, pdf_passwd=None):
        pages_count = None
        with open(pdf_path, "rb") as fp:
            try:
                pages_count = len(list(PDFPage.get_pages(fp, None, password=pdf_passwd, caching=False)))
            except Exception as e:
                print("!!! pages_count: {} error: {}".format(pdf_path, e))
        return pages_count

    @staticmethod
    def parse_text(pdf_path, pages: set() = None, pdf_passwd=None, codec="utf-8"):
        with open(pdf_path, "rb") as fp, io.StringIO() as output_string:
            rsrcmgr = PDFResourceManager(caching=True)
            device = TextConverter(rsrcmgr, output_string, codec=codec, laparams=LAParams())
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            for page in PDFPage.get_pages(fp, pages if pages is not None else None, password=pdf_passwd, caching=True):
                interpreter.process_page(page)
                print(output_string.getvalue())
            return output_string.getvalue()

    @staticmethod
    def parse_layout(pdf_path, pages: set() = None, pdf_passwd=None):
        layouts = dict()
        image_writer = ImageWriter(os.path.dirname(pdf_path))

        parser = PDFParser(open(pdf_path, 'rb'))
        document = PDFDocument(parser, pdf_passwd)
        if not document.is_extractable:
            raise PDFTextExtractionNotAllowed
        rsrc_mgr = PDFResourceManager()
        device_layout = PDFPageAggregator(rsrc_mgr, laparams=LAParams(detect_vertical=True, all_texts=True, ))
        interpreter = PDFPageInterpreter(rsrc_mgr, device_layout)
        # Make a page iterator
        pdf_pages = PDFPage.create_pages(document)
        for page_index, pdf_page in enumerate(pdf_pages, 1):
            if pages and page_index not in pages:
                continue
            # oh so stateful
            interpreter.process_page(pdf_page)
            layout = device_layout.get_result()
            layouts[page_index] = layout
        return layouts

    @staticmethod
    def extract_info(layouts, image_writer=None):
        images = dict()
        layouts_types_statistics = collections.Counter()
        for page_index, layout in layouts.items():
            # Look at all (nested) objects on each page
            for item in PdfminerWrapper.flat_iter(layout):
                layouts_types_statistics.update(type(item).__name__)
                if isinstance(item, LTImage):
                    if image_writer is not None:
                        image_writer.export_image(item)
                    if page_index not in images:
                        images[page_index] = list()
                    images[page_index].append(item)
        message('layouts_types_statistics:', ' '.join('{}:{}'.format(*tc) for tc in layouts_types_statistics.items()))
        return images


class PdfPlumberWrapper(object):
    def __init__(self, pdf_path, pdf_passwd=None):
        self.pdf_path = pdf_path
        self.pdf_obj = pdfplumber.open(pdf_path)

    def __del__(self):
        self.pdf_obj.close()

    def parse_pdf(self, pages: set() = None):
        page_data = dict()
        pdf = self.pdf_obj
        try:
            pdf_pages = dict()
            for index in range(len(pdf.pages)):
                page_number = pdf.pages[index].page_number
                if pages and page_number not in pages:
                    continue
                if page_number not in pdf_pages:
                    pdf_pages[page_number] = list()
                pdf_pages[page_number].append(pdf.pages[index])
            for page_number, page_pages in pdf_pages.items():
                if len(page_pages) != 1:
                    print("---> pdf {} page_number {} has {} pdf.Pages".format(self.pdf_path, page_number, len(page_pages)))
                texts, images = PdfPlumberWrapper.extract_info(page_pages)
                page_data[page_number] = (texts, images)
        except Exception as e:
            print("!!! {}".format(e))
        return page_data

    @staticmethod
    def extract_info(pdf_pages: list(), use_imagewriter=True):
        texts, images = "", list()
        for pdf_page in pdf_pages:
            # texts
            text = pdf_page.extract_text(x_tolerance=3, y_tolerance=3)
            texts += text if text else ""
            # images
            page_height = pdf_page.height
            for pdf_image in pdf_page.images:
                if use_imagewriter:
                    images.append(DictObject(pdf_image))
                else:
                    image_bbox = (pdf_image['x0'], page_height - pdf_image['y1'], pdf_image['x1'], page_height - pdf_image['y0'])
                    cropped_page = pdf_page.crop(image_bbox)
                    image_obj = cropped_page.to_image(resolution=400)
                    image_obj.save(bytes)
                    image_pil = pil.Image.open(io.BytesIO(bytes.getvalue()))
                    images.append(image_pil)
        return texts, images

    @staticmethod
    def record_images(images, output_dir_path, prefix=""):
        if not images:
            return None
        os.makedirs(output_dir_path, exist_ok=True)
        # when use_imagewriter = True
        for _, img_miner in enumerate(images):
            try:
                img_obj = DictObject(img_miner)
                img_obj.srcsize = (int(img_obj.srcsize[0]), int(img_obj.srcsize[1]))
                image_writer = ImageWriter(output_dir_path)
                image_name = image_writer.export_image(img_obj)
            except Exception as e:
                print("!!! record_images error: {}".format(e))
        return len(images)


class DocxWrapper(object):
    @staticmethod
    def extract_images(docx_path, output_dir=".", tmp_path=None):
        images_path = []
        docx_dir, docx_name = os.path.split(docx_path)
        if docx_path is None or output_dir is None:
            raise Exception("invalid arguments")
        tmp_path = docx_path if tmp_path is None else tmp_path
        tmp_path = os.path.join(tmp_path, docx_name)
        if not os.path.isabs(docx_path) or not os.path.isabs(tmp_path):
            tmp_path = os.path.abspath(tmp_path)
            docx_path = os.path.abspath(docx_path)
        if not os.path.exists(tmp_path) or not os.path.isdir(tmp_path):
            os.makedirs(tmp_path, 755);
        zip_path = os.path.join(tmp_path, docx_name + ".zip")

        shutil.copy(docx_path, zip_path)
        f = zipfile.ZipFile(zip_path, 'r')
        for file in f.namelist():
            f.extract(file, tmp_path)
        f.close()

        # 得到缓存文件夹中图片列表 并 清除临时文件
        pics = os.listdir(os.path.join(tmp_path, 'word/media'))
        for pic in pics:
            # 根据word的路径生成图片的名称
            new_path = os.path.join(output_dir, docx_name.replace(':', '') + '_' + pic)
            shutil.copy(os.path.join(tmp_path + '/word/media', pic), new_path)
            images_path.append(new_path)
        shutil.rmtree(tmp_path)
        return images_path

    @staticmethod
    def extract_texts_tables(docx_path):
        texts, images, tables = "", list(), dict()
        doc = docx.Document(docx_path)
        for inline_shape in doc.inline_shapes:
            images.append(inline_shape.type)
        for i, paragraph in enumerate(doc.paragraphs):
            texts += paragraph.text
            if i < len(doc.paragraphs) - 1:
                texts += "\n"
        # 单元格合并的特点
        #   每个单元格都有其对应的内存地址。单元格合并之后这几个合并的单元格会共享内存地址。
        #   不过，在枚举单元格时：
        #           按行枚举时，按行合并的单元格地址相同，按列合并的单元格地址还是不相同;
        #           按列枚举时，按列合并的单元格地址相同，按行合并的单元格地址依然不相同;
        #   所以，需要分别按行、列枚举，以合并行列上的单元格
        for i, table in enumerate(doc.tables):
            cell_merged_map = dict()
            cell_text = dict()
            for r, row in enumerate(table.rows):
                cell_text[r] = [cell.text for cell in row.cells]
                cell_merged_map[r] = [-1 for _ in row.cells]
            cell_addr = dict()
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    if cell not in cell_addr:
                        cell_addr[cell] = len(cell_addr)
                    cell_merged_map[r][c] = cell_addr[cell]
            cell_addr.clear()
            for c, col in enumerate(table.columns):
                for r, cell in enumerate(col.cells):
                    if cell not in cell_addr:
                        cell_addr[cell] = (r,c)
                    else:
                        merged_r, merged_c = cell_addr[cell]
                        cell_merged_map[r][c] = cell_merged_map[merged_r][merged_c]
            tables[i] = (cell_text, cell_merged_map)
        return texts, images, tables
