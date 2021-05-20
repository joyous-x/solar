# coding=utf-8

""" pdf-checker
pdf 相关工具的库主要有：at 2021-05-17
- [camelot](https://github.com/camelot-dev/camelot): 提取表格效果最好
    + 使用了 PyPDF2 来操作 pdf 文件
    + WARN：camelot\handlers.py: L73 打开文件后没有关闭，会句柄泄漏
- [pdfminer](https://github.com/pdfminer/pdfminer.six)
    + 可以解析 pdf 的树形结构，用于提取 文字、图片等
- [pdfplumber](https://github.com/jsvine/pdfplumber)
    + built on pdfminer.six

office 相关库：
- [python-docx](https://github.com/python-openxml/python-docx)
- [openpyxl](https://openpyxl.readthedocs.io/en/stable/)


# reference
#     https://openpyxl.readthedocs.io/en/stable/api/openpyxl.styles.fills.html
# note:
#     openpyxl 在插入一行时，如果后续有mergedcell，则会错乱：不能处理mergedcell的格式


mind:
- normal
    + NOTE
        - 不做 ocr，保证图片、文字、表格的位置、格式准确
    + 元素
        + image
            image 个数
            是否需要 ocr ？
        + table
            抽取 grid，计算 joint 的 行列 分布
            抽取 text 计算 每行、每列 的文本相似度
        + text
            pdf 读取 和 doc 读取，文本相似度计算
            格式: 可以忽略
- scanned ：ocr (baseline ?)
    + image
    + table
    + text

todo:
建议.
1. 增加 pdf 格式检查，非法格式不操作
2. 超过 60 页的 pdf 文件，耗时非常久(会超时失败，当然，具体耗时跟pdf内容也有关系)，，可以增加分拆动作
3. 语言：中、英、俄、韩



"""
import os
import io
import shutil
import numpy as np
import PIL as pil

import zipfile
import docx

import camelot
import PyPDF2
try:
    import pdfplumber
except Exception as e:
    pdfplumber = None

from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage, PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator, TextConverter
from pdfminer.layout import LAParams, LTContainer,LTImage
from pdfminer.image import ImageWriter
import pdfminer.high_level
import pdfminer.layout

import collections

from openpyxl import load_workbook, Workbook
from openpyxl.worksheet.table import Table, TableStyleInfo


def msg(*args, **kwargs):
    print(' '.join(map(str, args)), **kwargs)  # noqa E999


class CamelotWrapper(object):
    """ parse tables in pdf with Camelot

    reference: https://camelot-py.readthedocs.io/en/master/user/advanced.html
    """
    def __init__(self, filepath, passwd=None, pages="all", line_scale=40):
        """
        pages : str, optional (default: '1')
            Comma-separated page numbers.
            Example: '1,3,4' or '1,4-end' or 'all'.
        """
        self.filepath, self.passwd = filepath, passwd
        self.line_scale = line_scale
        self.pages = pages
        self.tables = None
        try:
            self.tables = camelot.read_pdf(filepath, pages="all" if not pages else pages, password=passwd, line_scale=line_scale)
        except Exception as e:
            print("!!! {} error: {}".format(filepath, e))

    def parse_tables(self, pages: set() = None, accuracy_min=10, debug=False):
        if self.tables is None:
            return None
        if debug:
            for table in self.tables:
                data_frame = table.df
                parsing_report = table.parsing_report
                print(parsing_report, "\n", data_frame)
        tables = list(filter(lambda t: t.parsing_report['accuracy'] > accuracy_min, self.tables))
        if pages is not None:
            tables = list(filter(lambda t: t.parsing_report['page'] in pages, tables))
        return tables

    @staticmethod
    def show(table, kind='grid'):
        camelot.plot(table, kind=kind).show()

    @staticmethod
    def pages_count(file_path, pdf_passwd=None):
        pages_count = None
        try:
            pdf_obj = PyPDF2.PdfFileReader(file_path, strict=False)
            pages_count = pdf_obj.getNumPages()
        except Exception as e:
            print("!!! {} error: {}".format(file_path, e))
        return pages_count

    @staticmethod
    def split_pages(file_path, output_dir=None, pdf_passwd=None, debug=False):
        print("==> ready to split : {}".format(file_path))
        file_dir, file_name = os.path.split(file_path)
        file_name, file_suffix = os.path.splitext(file_name)
        if not output_dir:
            output_dir = file_dir
        pdf_obj = PyPDF2.PdfFileReader(file_path, strict=False)
        for pageno in range(1, pdf_obj.getNumPages()+1):
            pdf_writer = PyPDF2.PdfFileWriter()
            pdf_writer.addPage(pdf_obj.getPage(pageno-1))
            with open(os.path.join(output_dir, "{}.{}{}".format(file_name,pageno,file_suffix)), 'wb') as output_pdf:
                pdf_writer.write(output_pdf)
        print("===> end: total pages {}".format(pdf_obj.getNumPages()))
        return pdf_obj.getNumPages()

    @staticmethod
    def record_tables(tables, output_dir_path, prefix=""):
        if not tables:
            return None
        tables_cells_bound= list()
        for table in tables:
            table_cells_bound = list()
            for r in range(table.shape[0]):
                cells = [table.cells[r][c] for c in range(table.shape[1])]
                table_cells_bound.append([(cell.bottom, cell.top, cell.left, cell.right) for cell in cells])
            tables_cells_bound.append(table_cells_bound)
        for index, table in enumerate(tables_cells_bound):
            lines = [",".join([str(t) for t in line]) for line in table]
            with open(os.path.join(output_dir_path, "{}table_{}_bound.txt".format(prefix, index)), "a+") as f:
                f.write("\n".join(lines))
        for table in tables:
            lines = [" | ".join(line) for line in table.data]
            with open(os.path.join(output_dir_path, "{}table_{}_texts.txt".format(prefix, index)), "a+") as f:
                f.write("\n".join(lines))
        return len(tables)


class PdfminerWrapper(object):
    """
    reference: https://pdfminersix.readthedocs.io/en/latest/index.html
    """
    @staticmethod
    def flat_iter(obj):
        yield obj
        if isinstance(obj, LTContainer):
            for ob in obj:
                yield from PdfminerWrapper.flat_iter(ob)

    @staticmethod
    def pages_count(pdf_path, pdf_passwd=None):
        pages_count = None
        with open(pdf_path, "rb") as fp:
            try:
                pages_count = len(list(PDFPage.get_pages(fp, None, password=pdf_passwd, caching=False)))
            except Exception as e:
                print("!!! pages_count: {} error: {}".format(pdf_path, e))
        return pages_count

    @staticmethod
    def parse_text(pdf_path, pages: set() = None, pdf_passwd=None, codec="utf-8"):
        with open(pdf_path, "rb") as fp, io.StringIO() as output_string:
            rsrcmgr = PDFResourceManager(caching=True)
            device = TextConverter(rsrcmgr, output_string, codec=codec, laparams=LAParams())
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            for page in PDFPage.get_pages(fp, pages if pages is not None else None, password=pdf_passwd, caching=True):
                interpreter.process_page(page)
                print(output_string.getvalue())
            return output_string.getvalue()

    @staticmethod
    def parse_layout(pdf_path, pages: set() = None, pdf_passwd=None):
        layouts = dict()
        image_writer = ImageWriter(os.path.dirname(pdf_path))

        parser = PDFParser(open(pdf_path, 'rb'))
        document = PDFDocument(parser, pdf_passwd)
        if not document.is_extractable:
            raise PDFTextExtractionNotAllowed
        rsrc_mgr = PDFResourceManager()
        device_layout = PDFPageAggregator(rsrc_mgr, laparams=LAParams(detect_vertical=True, all_texts=True, ))
        interpreter = PDFPageInterpreter(rsrc_mgr, device_layout)
        # Make a page iterator
        pdf_pages = PDFPage.create_pages(document)
        for page_index, pdf_page in enumerate(pdf_pages, 1):
            if pages and page_index not in pages:
                continue
            # oh so stateful
            interpreter.process_page(pdf_page)
            layout = device_layout.get_result()
            layouts[page_index] = layout
        return layouts

    @staticmethod
    def extract_info(layouts, image_writer=None):
        images = dict()
        layouts_types_statistics = collections.Counter()
        for page_index, layout in layouts.items():
            # Look at all (nested) objects on each page
            for item in PdfminerWrapper.flat_iter(layout):
                layouts_types_statistics.update(type(item).__name__)
                if isinstance(item, LTImage):
                    if image_writer is not None:
                        image_writer.export_image(item)
                    if page_index not in images:
                        images[page_index] = list()
                    images[page_index].append(item)
        msg('layouts_types_statistics:', ' '.join('{}:{}'.format(*tc) for tc in layouts_types_statistics.items()))
        return images


class DictObject(dict):
    def __getattr__(self, key):
        return self.get(key)

    def __setattr__(self, key, value):
        self[key] = value


class PdfPlumberWrapper(object):
    def __init__(self, pdf_path, pdf_passwd=None):
        self.pdf_path = pdf_path
        self.pdf_obj = pdfplumber.open(pdf_path)

    def __del__(self):
        self.pdf_obj.close()

    def parse_pdf(self, pages: set() = None):
        page_data = dict()
        pdf = self.pdf_obj
        try:
            pdf_pages = dict()
            for index in range(len(pdf.pages)):
                page_number = pdf.pages[index].page_number
                if pages and page_number not in pages:
                    continue
                if page_number not in pdf_pages:
                    pdf_pages[page_number] = list()
                pdf_pages[page_number].append(pdf.pages[index])
            for page_number, page_pages in pdf_pages.items():
                if len(page_pages) != 1:
                    print("---> pdf {} page_number {} has {} pdf.Pages".format(self.pdf_path, page_number, len(page_pages)))
                texts, images = PdfPlumberWrapper.extract_info(page_pages)
                page_data[page_number] = (texts, images)
        except Exception as e:
            print("!!! {}".format(e))
        return page_data

    @staticmethod
    def extract_info(pdf_pages: list(), use_imagewriter=True):
        texts, images = "", list()
        for pdf_page in pdf_pages:
            # texts
            text = pdf_page.extract_text(x_tolerance=3, y_tolerance=3)
            texts += text if text else ""
            # images
            page_height = pdf_page.height
            for pdf_image in pdf_page.images:
                if use_imagewriter:
                    images.append(DictObject(pdf_image))
                else:
                    image_bbox = (pdf_image['x0'], page_height - pdf_image['y1'], pdf_image['x1'], page_height - pdf_image['y0'])
                    cropped_page = pdf_page.crop(image_bbox)
                    image_obj = cropped_page.to_image(resolution=400)
                    image_obj.save(bytes)
                    image_pil = pil.Image.open(io.BytesIO(bytes.getvalue()))
                    images.append(image_pil)
        return texts, images

    @staticmethod
    def record_images(images, output_dir_path, prefix=""):
        if not images:
            return None
        os.makedirs(output_dir_path, exist_ok=True)
        # when use_imagewriter = True
        for _, img_miner in enumerate(images):
            try:
                img_obj = DictObject(img_miner)
                img_obj.srcsize = (int(img_obj.srcsize[0]), int(img_obj.srcsize[1]))
                image_writer = ImageWriter(output_dir_path)
                image_name = image_writer.export_image(img_obj)
            except Exception as e:
                print("!!! record_images error: {}".format(e))
        return len(images)


class DocxWrapper(object):
    @staticmethod
    def extract_images(docx_path, output_dir=".", tmp_path=None):
        images_path = []
        docx_dir, docx_name = os.path.split(docx_path)
        if docx_path is None or output_dir is None:
            raise Exception("invalid arguments")
        tmp_path = docx_path if tmp_path is None else tmp_path
        tmp_path = os.path.join(tmp_path, docx_name)
        if not os.path.isabs(docx_path) or not os.path.isabs(tmp_path):
            tmp_path = os.path.abspath(tmp_path)
            docx_path = os.path.abspath(docx_path)
        if not os.path.exists(tmp_path) or not os.path.isdir(tmp_path):
            os.makedirs(tmp_path, 755);
        zip_path = os.path.join(tmp_path, docx_name + ".zip")

        shutil.copy(docx_path, zip_path)
        f = zipfile.ZipFile(zip_path, 'r')
        for file in f.namelist():
            f.extract(file, tmp_path)
        f.close()

        # 得到缓存文件夹中图片列表 并 清除临时文件
        pics = os.listdir(os.path.join(tmp_path, 'word/media'))
        for pic in pics:
            # 根据word的路径生成图片的名称
            new_path = os.path.join(output_dir, docx_name.replace(':', '') + '_' + pic)
            shutil.copy(os.path.join(tmp_path + '/word/media', pic), new_path)
            images_path.append(new_path)
        shutil.rmtree(tmp_path)
        return images_path

    @staticmethod
    def extract_texts_tables(docx_path):
        texts, images, tables = "", list(), dict()
        doc = docx.Document(docx_path)
        for inline_shape in doc.inline_shapes:
            images.append(inline_shape.type)
        for i, paragraph in enumerate(doc.paragraphs):
            texts += paragraph.text
            if i < len(doc.paragraphs) - 1:
                texts += "\n"
        # 单元格合并的特点
        #   每个单元格都有其对应的内存地址。单元格合并之后这几个合并的单元格会共享内存地址。
        #   不过，在枚举单元格时：
        #           按行枚举时，按行合并的单元格地址相同，按列合并的单元格地址还是不相同;
        #           按列枚举时，按列合并的单元格地址相同，按行合并的单元格地址依然不相同;
        #   所以，需要分别按行、列枚举，以合并行列上的单元格
        for i, table in enumerate(doc.tables):
            cell_merged_map = dict()
            cell_text = dict()
            for r, row in enumerate(table.rows):
                cell_text[r] = [cell.text for cell in row.cells]
                cell_merged_map[r] = [-1 for _ in row.cells]
            cell_addr = dict()
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    if cell not in cell_addr:
                        cell_addr[cell] = len(cell_addr)
                    cell_merged_map[r][c] = cell_addr[cell]
            cell_addr.clear()
            for c, col in enumerate(table.columns):
                for r, cell in enumerate(col.cells):
                    if cell not in cell_addr:
                        cell_addr[cell] = (r,c)
                    else:
                        merged_r, merged_c = cell_addr[cell]
                        cell_merged_map[r][c] = cell_merged_map[merged_r][merged_c]
            tables[i] = (cell_text, cell_merged_map)
        return texts, images, tables


class DocChecker(object):
    def __init__(self, filepath, passwd=None):
        self.filepath, self.passwd = filepath, passwd
        self.plumber = PdfPlumberWrapper(filepath, passwd)

    @staticmethod
    def walk_dir(file_dir, file_name):
        files = []
        if file_name is not None:
            files.append(os.path.join(file_dir, file_name))
        else:
            for root, dirs, names in os.walk(file_dir, topdown=False):
                for sub_name in names:
                    files.append(os.path.join(root, sub_name))
                for sub_dir in dirs:
                    for ss_root, _, ss_files in os.walk(os.path.join(root, sub_dir), topdown=False):
                        for ss_name in files:
                            files.append(os.path.join(ss_root, ss_name))
        return files

    def pdf_layouts(self):
        """ 获取布局中的元素信息
        """
        pass

    def parse_pdf(self, pages: set() = None, writer=None, output_dir_root=None):
        """ 获取 pdf 中的 texts, images, tables
        注意：为了优化 tables 的检查时间较长的问题，先判断是否有文本，如果没有文本检出，则认为没有表格
        """
        texts, images, tables = dict(), dict(), dict()
        for page_number in pages:
            texts[page_number], images[page_number], tables[page_number] = None, None, None
            page_data = self.plumber.parse_pdf({page_number})
            for _, texts_images in page_data.items():
                texts[page_number] = texts_images[0]
                images[page_number] = texts_images[1]
            if page_number in texts and texts[page_number] and len(texts[page_number]) > 0:
                pdf_tables = CamelotWrapper(self.filepath, self.passwd, pages=str(page_number))
                tables[page_number] = pdf_tables.parse_tables({page_number})
            else:
                tables[page_number] = None
            if writer is not None:
                writer(self.filepath, page_number, texts[page_number], images[page_number], tables[page_number], output_dir_root, do_record=False)
        return texts, tables, images

    def parse_docx(self, docx_path, output_dir="."):
        texts, images, tables = DocxWrapper.extract_texts_tables(docx_path)
        images = DocxWrapper.extract_images(docx_path, output_dir, "./tmp")
        return texts, images, tables

    def parse_xls(self):
        templateDir = os.path.dirname(os.path.abspath(__file__))
        wb = load_workbook(os.path.join(templateDir, "bill_template.xlsx"))
        shPreface = wb.get_sheet_by_name(wb.sheetnames[0])
        shDetail = wb.get_sheet_by_name(wb.sheetnames[1])

        texts, grids = None, None
        pass
        return texts, grids

    def pdf_images_need_ocr(self):
        pass

    @staticmethod
    def record_page_info(filepath, pageno, texts, images, tables, output_dir_root, do_record=True):
        file_dir, file_name = os.path.split(filepath)
        file_name, file_suffix = os.path.splitext(file_name)

        text_sample, images_count, table_count, text_count = "", 0, 0, 0
        if texts:
            text_count += len(texts.replace(' ', ''))
            text_sample = texts[:min(text_count, 30)]
        images_count += len(images) if images else 0
        table_count += len(tables) if tables else 0
        output_dir_name = "{}{}.{}.{}_{}_{}".format(file_name, file_suffix, pageno, images_count, table_count, text_count)
        output_dir_path = os.path.join(output_dir_root, output_dir_name)

        if do_record:
            shutil.rmtree(output_dir_path, ignore_errors=True)
            os.makedirs(output_dir_path, exist_ok=True)

            with open(os.path.join(output_dir_path, "texts.txt"), "a+", encoding='utf-8') as f:
                f.write(texts.replace(' ', ''))
            if images:
                PdfPlumberWrapper.record_images(images, os.path.join(output_dir_path, "images"))
            if tables:
                CamelotWrapper.record_tables(tables, output_dir_path)
        print("\t page({}) has {} images {} tables, texts: {} ".format(pageno, images_count, table_count, repr(text_sample)))
        return text_count, images_count, table_count

    @staticmethod
    def extract_pdf(file_path, target_dir, prefix=""):
        text_count, images_count, table_count = 0, 0, 0
        pages_count = CamelotWrapper.pages_count(file_path)
        if pages_count is None or pages_count > 100:
            return pages_count if pages_count else -1, -1, -1, -1
        print("==> prepare: {} , total pages: {}".format(file_path, pages_count))
        checker = DocChecker(file_path, None)
        for index in range(1, 1+pages_count):
            _text, _table, _images = checker.parse_pdf({index}, writer=DocChecker.record_page_info, output_dir_root=target_dir)
            text_count += len(_text[index]) if _text[index] else 0
            images_count += len(_images[index]) if _images[index] else 0
            table_count += len(_table[index]) if _table[index] else 0
        print("===> {} has {} words {} images {} tables".format(file_path, text_count, images_count, table_count))
        return pages_count, text_count, images_count, table_count


def __run_tests(pdf_dir, pdf_type, output_dir):
    files = DocChecker.walk_dir(os.path.join(pdf_dir, pdf_type), None)
    for file_path in files:
        result = DocChecker.extract_pdf(file_path, output_dir, "word_")
        if result:
            page_count, text_count, images_count, table_count = result
            statistics_file = os.path.join(os.path.dirname(pdf_dir), pdf_type + ".txt")
            with open(statistics_file, "a+") as f:
                f.write("{}, {}, {}, {}, {} \n".format(file_path, page_count, text_count, images_count, table_count))
            target_dir = os.path.join(os.path.dirname(pdf_dir), "processed", pdf_type)
            if not os.path.exists(target_dir) or not os.path.isdir(target_dir):
                os.makedirs(target_dir, exist_ok=True)
            _, file_name = os.path.split(file_path)
            shutil.move(file_path, os.path.join(target_dir, file_name))


if __name__ == "__main__":
    __run_tests("D:\\Workspace\\datasets\\pdf_convert\\raw", "pdf_to_ppt", "D:\\Workspace\\datasets\\pdf_convert\\output")
